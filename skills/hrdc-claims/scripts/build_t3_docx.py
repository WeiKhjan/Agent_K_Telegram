#!/usr/bin/env python3
"""
HRDC PSMB/SBL-KHAS/T3/01 Attendance List — DOCX Template Filler

Opens the master .docx template (downloaded from Google Docs), fills in
the attendance table with trainee data, and saves a new .docx file.

The template has two attendance tables (Table 2 = Day 1, Table 7 = Day 2).
Each table has a header row + 10 empty rows. If more rows are needed,
extra rows are cloned from the template.

For multi-day training, generates one file with all days filled.

Usage:
  python3 build_t3_docx.py \
    --course "AI Workflow Automation" \
    --dates "2/4/2026,3/4/2026" \
    --participants '[{"name":"JOHN DOE","employer":"ABC SDN BHD","nric":"900101011234"}]' \
    --certifier-name "Goh Hen Yee" \
    --certifier-designation "Director" \
    --cert-date "6/4/2026" \
    --output "~/PSMB_SBL_KHAS_T3_01_FILLED.docx"

  OR with Excel input:
  python3 build_t3_docx.py \
    --course "AI Workflow Automation" \
    --dates "2/4/2026,3/4/2026" \
    --excel "path/to/file.xlsx" \
    --sheet "Name List" \
    --certifier-name "Goh Hen Yee" \
    --output "~/PSMB_SBL_KHAS_T3_01_FILLED.docx"
"""

import argparse
import json
import os
import sys
from copy import deepcopy

from docx import Document
from docx.shared import Pt

# ── Template path ────────────────────────────────────────────────────────────
TEMPLATE_PATH = os.path.expanduser(
    "~/Agent_K_Telegram/templates/PSMB_SBL_KHAS_T3_01_template.docx"
)

# Template structure: tables[2] and tables[7] are the attendance tables
# for Day 1 and Day 2 respectively. Each has header row + 10 empty data rows.
# Tables[1]/[6] = course info (date in row 1, cell 2)
# Tables[3]/[8] = signature block (certifier info)
DAY_TABLE_INDICES = [2, 7]       # attendance tables
DAY_INFO_INDICES = [1, 6]        # course info tables
DAY_SIG_INDICES = [3, 8]         # signature tables

ROWS_PER_TEMPLATE_PAGE = 10     # empty rows in template per table


def parse_args():
    p = argparse.ArgumentParser(description="Fill HRDC T3 Attendance List (.docx)")
    p.add_argument("--course", required=True, help="Course title")
    p.add_argument("--dates", required=True, help="Comma-separated training dates, e.g. 2/4/2026,3/4/2026")
    p.add_argument("--participants", help="JSON array of participant dicts (name, employer, nric)")
    p.add_argument("--excel", help="Path to Excel file with participant data")
    p.add_argument("--sheet", default="Name List", help="Sheet name in Excel (default: Name List)")
    p.add_argument("--certifier-name", default="Goh Hen Yee")
    p.add_argument("--certifier-designation", default="Director")
    p.add_argument("--cert-date", help="Certification date (default: last training date)")
    p.add_argument("--output", required=True, help="Output .docx path")
    p.add_argument("--template", default=TEMPLATE_PATH, help="Template .docx path")
    return p.parse_args()


def get_sex(nric: str) -> str:
    """Determine sex from NRIC last digit: odd = M, even = F."""
    nric_clean = nric.replace("-", "").strip()
    if not nric_clean:
        return ""
    return "M" if int(nric_clean[-1]) % 2 == 1 else "F"


def load_from_excel(excel_path: str, sheet_name: str = "Name List") -> list:
    """Load participant data from Excel Name List tab.

    Expected columns: A=Company Name, B=Name, C=NRIC
    """
    import openpyxl
    wb = openpyxl.load_workbook(excel_path)
    ws = wb[sheet_name]

    participants = []
    for row in ws.iter_rows(min_row=2, max_row=ws.max_row, values_only=False):
        company = row[0].value  # Column A
        name = row[1].value     # Column B
        nric = row[2].value     # Column C

        if not name:
            continue

        nric_str = str(nric) if nric else ""
        # Pad NRIC to 12 digits if numeric
        if nric_str and nric_str.isdigit() and len(nric_str) < 12:
            nric_str = nric_str.zfill(12)

        participants.append({
            "name": str(name).upper().strip(),
            "employer": str(company).upper().strip() if company else "",
            "nric": nric_str,
        })

    return participants


def fill_cell(cell, text):
    """Fill a cell with text, preserving existing paragraph/run formatting."""
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.text = ""
    if cell.paragraphs:
        p = cell.paragraphs[0]
        if p.runs:
            p.runs[0].text = text
        else:
            p.add_run(text)


def clone_row(table, template_row_idx: int):
    """Clone a row from the table XML and append it."""
    template_tr = table.rows[template_row_idx]._tr
    new_tr = deepcopy(template_tr)
    table._tbl.append(new_tr)
    return table.rows[-1]


def fill_attendance_table(table, participants: list):
    """Fill an attendance table with participant data, adding rows as needed."""
    num_existing = len(table.rows) - 1  # minus header row
    num_needed = len(participants)

    # Add extra rows if needed
    extra = num_needed - num_existing
    if extra > 0:
        for _ in range(extra):
            clone_row(table, 1)  # clone formatting from first data row

    # Fill data rows
    for idx, p in enumerate(participants):
        row = table.rows[idx + 1]  # skip header
        nric = p.get("nric", "")
        sex = p.get("sex", "") or get_sex(nric)
        citizenship = p.get("citizenship", "MALAYSIAN" if nric else "")

        fill_cell(row.cells[0], str(idx + 1))          # No.
        fill_cell(row.cells[1], p.get("name", ""))      # Name of Trainee
        fill_cell(row.cells[2], p.get("employer", ""))   # Name of Employer
        fill_cell(row.cells[3], nric)                    # NRIC
        fill_cell(row.cells[4], citizenship)             # Citizenship
        fill_cell(row.cells[5], sex)                     # Sex
        # cells[6] = Signature (leave blank)

    # Clear any leftover empty rows beyond the data
    for idx in range(num_needed, num_existing):
        row = table.rows[idx + 1]
        for cell in row.cells:
            fill_cell(cell, "")


def update_date(info_table, date_str: str):
    """Update the training date in the course info table."""
    # Row 1, Cell 2 = date
    if len(info_table.rows) > 1 and len(info_table.rows[1].cells) > 2:
        fill_cell(info_table.rows[1].cells[2], date_str)


def update_certifier(sig_table, name: str, designation: str, cert_date: str):
    """Update certifier info in the signature block table."""
    # Row 0: NAME : [name] ... SIGNATURE : ...
    if len(sig_table.rows) > 0 and len(sig_table.rows[0].cells) > 2:
        fill_cell(sig_table.rows[0].cells[2], name)
    # Row 1: DESIGNATION : [designation] ... DATE : [date]
    if len(sig_table.rows) > 1:
        row = sig_table.rows[1]
        if len(row.cells) > 2:
            fill_cell(row.cells[2], designation)
        if len(row.cells) > 6:
            fill_cell(row.cells[6], cert_date)


def build_docx(args):
    output_path = os.path.expanduser(args.output)
    os.makedirs(os.path.dirname(output_path) or ".", exist_ok=True)

    template_path = os.path.expanduser(args.template)
    if not os.path.exists(template_path):
        print(f"❌ Template not found: {template_path}")
        sys.exit(1)

    # Load participants
    if args.excel:
        participants = load_from_excel(args.excel, args.sheet)
    elif args.participants:
        participants = json.loads(args.participants)
        # Uppercase names
        for p in participants:
            p["name"] = p.get("name", "").upper()
            p["employer"] = p.get("employer", "").upper()
    else:
        print("❌ Either --participants or --excel is required")
        sys.exit(1)

    dates = [d.strip() for d in args.dates.split(",")]
    cert_date = args.cert_date or dates[-1]

    # Template supports 2 days. For more days, we'd need to duplicate sections.
    # Current template has exactly 2 day sections.
    if len(dates) > 2:
        print(f"⚠️  Template supports max 2 days. Using first 2 of {len(dates)} dates.")
        dates = dates[:2]

    doc = Document(template_path)

    for day_idx, date_str in enumerate(dates):
        if day_idx >= len(DAY_TABLE_INDICES):
            break

        # Fill attendance table
        att_table = doc.tables[DAY_TABLE_INDICES[day_idx]]
        fill_attendance_table(att_table, participants)

        # Update date
        info_table = doc.tables[DAY_INFO_INDICES[day_idx]]
        update_date(info_table, date_str)

        # Update certifier
        sig_table = doc.tables[DAY_SIG_INDICES[day_idx]]
        update_certifier(sig_table, args.certifier_name,
                        args.certifier_designation, cert_date)

    doc.save(output_path)
    print(f"✅ T3 saved: {output_path}  ({len(dates)} day(s), {len(participants)} participants)")


if __name__ == "__main__":
    args = parse_args()
    build_docx(args)
