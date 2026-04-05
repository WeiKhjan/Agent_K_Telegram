#!/usr/bin/env python3
"""Update HRDC Confirmation Letter template with provided details.

Usage:
    python update_confirmation.py \
        --date "26th March 2026" \
        --company "Daxin KF&C PLT" \
        --address "54, Jln Kempas Utama 2/2, Taman Kempas Utama,\n81200 Johor Bahru, Johor Darul Ta'zim" \
        --subject "CONFIRMATION ON AI AGENTIC AUTOMATION WITH N8N ON 2nd AND 3rd APRIL 2026" \
        --workshop "AI Agentic Automation with n8n" \
        --participants 1 \
        --day1-date "2nd April 2026" \
        --day2-date "3rd April 2026" \
        --fees "3,780.00" \
        --output "/path/to/output.docx"

Optional:
    --template  Path to template (default: skill asset)
    --location  Override location (default: keep template value)
"""

import argparse
import os
import re
from docx import Document


TEMPLATE_PATH = os.path.join(
    os.path.dirname(os.path.dirname(os.path.abspath(__file__))),
    "assets", "confirmation_letter_template.docx"
)


def set_runs_text(runs, new_text):
    """Put all text in first run, clear the rest. Preserves formatting of first run."""
    if not runs:
        return
    runs[0].text = new_text
    for r in runs[1:]:
        r.text = ""


def update_date_runs(runs, new_date_str):
    """Update date runs that are split like: 'day','th',' ','Month',' ','202','6'.

    new_date_str should be like '26th March 2026'.
    Handles both 6-run and 7-run variants.
    """
    m = re.match(r"(\d+)(st|nd|rd|th)\s+(\w+)\s+(\d{4})", new_date_str)
    if not m:
        set_runs_text(runs, new_date_str)
        return

    day, suffix, month, year = m.groups()

    if len(runs) >= 6:
        runs[0].text = day
        runs[1].text = suffix
        runs[2].text = " "
        runs[3].text = month
        # Handle both 'Month ' (single run) and 'Month',' ' (two runs) patterns
        if len(runs) == 6:
            runs[3].text = month + " "
            runs[4].text = year[:3]
            runs[5].text = year[3:]
        else:  # 7+ runs
            runs[3].text = month
            runs[4].text = " "
            runs[5].text = year[:3]
            runs[6].text = year[3:]
        for r in runs[max(6, 7):]:
            r.text = ""
    else:
        set_runs_text(runs, new_date_str)


def update_fees_runs(runs, new_fees):
    """Update fee runs that are split digit-by-digit.

    new_fees should be like '3,780.00' (without RM prefix).
    """
    if len(runs) >= 2:
        runs[1].text = new_fees + "."
        for r in runs[2:]:
            if r.text == "\n":
                continue
            r.text = ""
    else:
        set_runs_text(runs, f"Our confirmed fees are RM {new_fees}.\n")


def find_paragraph(doc, keyword):
    """Find paragraph containing keyword, return (index, paragraph)."""
    for i, p in enumerate(doc.paragraphs):
        if keyword in p.text:
            return i, p
    return None, None


def main():
    parser = argparse.ArgumentParser(description="Update HRDC Confirmation Letter")
    parser.add_argument("--date", required=True, help="Letter date, e.g. '26th March 2026'")
    parser.add_argument("--company", required=True, help="Company name")
    parser.add_argument("--address", required=True, help="Company address (use \\n for line breaks)")
    parser.add_argument("--subject", required=True, help="Subject line (all caps)")
    parser.add_argument("--workshop", help="Workshop name (for table)")
    parser.add_argument("--participants", type=int, required=True, help="Number of participants")
    parser.add_argument("--day1-date", required=True, help="Day 1 date, e.g. '2nd April 2026'")
    parser.add_argument("--day2-date", required=True, help="Day 2 date, e.g. '3rd April 2026'")
    parser.add_argument("--fees", required=True, help="Confirmed fees amount, e.g. '3,780.00'")
    parser.add_argument("--location", help="Override training location")
    parser.add_argument("--template", default=TEMPLATE_PATH, help="Path to template docx")
    parser.add_argument("--output", required=True, help="Output file path")
    args = parser.parse_args()

    doc = Document(args.template)

    # --- Date: find paragraph containing "Date:" ---
    _, p_date = find_paragraph(doc, "Date:")
    if p_date:
        date_runs = p_date.runs[1:]  # skip 'Date: ' prefix
        update_date_runs(date_runs, args.date)

    # --- Company + Address: find paragraph after "Board of Directors" ---
    board_idx, p_board = find_paragraph(doc, "Board of Directors")
    if p_board:
        # Company paragraph is 2 after Board of Directors (skip blank)
        for offset in range(1, 4):
            p_company = doc.paragraphs[board_idx + offset]
            if p_company.text.strip():
                break
        address_text = args.address.replace("\\n", "\n")
        full_text = f"{args.company}\n{address_text}\n"
        # Try to fit into existing run structure
        if len(p_company.runs) >= 5:
            p_company.runs[0].text = args.company
            # Clear middle runs, use one for address
            for r in p_company.runs[1:]:
                r.text = ""
            p_company.runs[1].text = "\n"
            address_lines = address_text.split("\n")
            if len(p_company.runs) >= 6 and len(address_lines) >= 2:
                p_company.runs[2].text = address_lines[0]
                p_company.runs[3].text = "\n"
                p_company.runs[4].text = address_lines[1]
                p_company.runs[5].text = "\n" if len(p_company.runs) > 5 else ""
                for r in p_company.runs[6:]:
                    r.text = ""
            else:
                p_company.runs[2].text = address_text + "\n"
        else:
            set_runs_text(p_company.runs, full_text)

    # --- Subject: find paragraph containing "CONFIRMATION" ---
    _, p_subject = find_paragraph(doc, "CONFIRMATION")
    if p_subject:
        set_runs_text(p_subject.runs, args.subject)

    # --- Fees: find paragraph containing "confirmed fees" ---
    _, p_fees = find_paragraph(doc, "confirmed fees")
    if p_fees:
        update_fees_runs(p_fees.runs, args.fees)

    # --- Table 0: Workshop details ---
    table0 = doc.tables[0]

    # Workshop name (row 0, cell 2)
    if args.workshop:
        cell = table0.rows[0].cells[2]
        set_runs_text(cell.paragraphs[0].runs, args.workshop)

    # Participants (row 1, cell 2)
    cell = table0.rows[1].cells[2]
    p = cell.paragraphs[0]
    if len(p.runs) >= 3:
        p.runs[0].text = str(args.participants)
    else:
        set_runs_text(p.runs, f"{args.participants} pax")

    # Location (row 2, cell 2)
    if args.location:
        cell = table0.rows[2].cells[2]
        set_runs_text(cell.paragraphs[0].runs, args.location)

    # --- Table 1: Day 1 agenda date ---
    t1_cell = doc.tables[1].rows[1].cells[0]
    for p in t1_cell.paragraphs:
        if p.runs and any(c.isdigit() for c in p.text) and "am" not in p.text.lower():
            update_date_runs(p.runs, args.day1_date)

    # --- Table 2: Day 2 agenda date ---
    t2_cell = doc.tables[2].rows[1].cells[0]
    for p in t2_cell.paragraphs:
        if p.runs and any(c.isdigit() for c in p.text) and "am" not in p.text.lower():
            update_date_runs(p.runs, args.day2_date)

    # Save
    os.makedirs(os.path.dirname(os.path.abspath(args.output)), exist_ok=True)
    doc.save(args.output)
    print(f"Saved: {args.output}")


if __name__ == "__main__":
    main()
